from flask import Flask, request, jsonify, render_template_string
from flask_cors import CORS
import os
from groq import Groq
import pypdf
import docx
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def extract_text_from_pdf(file):
    reader = pypdf.PdfReader(file)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

def extract_text_from_txt(file):
    return file.read().decode('utf-8', errors='ignore')

def analyze_resume(text):
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": """You are an expert resume reviewer and career coach. 
                Analyze the resume and provide detailed feedback in the following format:

                ## Overall Score
                Give a score out of 10 with brief explanation.

                ## Strengths
                List 3-5 strong points of the resume.

                ## Areas for Improvement
                List 3-5 specific areas that need improvement.

                ## Skills Analysis
                Analyze technical and soft skills present and missing.

                ## Structure & Formatting
                Comment on the resume structure, layout, and formatting.

                ## Specific Suggestions
                Give 5 actionable specific suggestions to improve the resume.

                ## Keywords Missing
                List important keywords that should be added based on the content.

                Be specific, constructive, and helpful."""
            },
            {
                "role": "user",
                "content": f"Please analyze this resume:\n\n{text}"
            }
        ]
    )
    return response.choices[0].message.content

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Resume Analyzer</title>
    <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Plus Jakarta Sans', sans-serif; background: #f0faf4; min-height: 100vh; }
        .navbar { background: white; padding: 16px 40px; display: flex; align-items: center; gap: 12px; box-shadow: 0 2px 12px rgba(0,0,0,0.06); }
        .logo { width: 40px; height: 40px; background: linear-gradient(135deg, #2ecc71, #27ae60); border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 18px; }
        .navbar h1 { font-size: 20px; font-weight: 800; color: #222; }
        .container { max-width: 900px; margin: 40px auto; padding: 0 20px; }
        .hero { text-align: center; margin-bottom: 40px; }
        .hero h2 { font-size: 36px; font-weight: 800; color: #1a1a1a; margin-bottom: 12px; }
        .hero p { font-size: 16px; color: #888; font-weight: 500; }
        .upload-card { background: white; border-radius: 24px; padding: 40px; box-shadow: 0 4px 24px rgba(0,0,0,0.07); margin-bottom: 32px; }
        .upload-area { border: 2px dashed #d4f5e4; border-radius: 16px; padding: 48px; text-align: center; cursor: pointer; transition: all 0.2s; background: #f9fffe; }
        .upload-area:hover { border-color: #2ecc71; background: #f0faf4; }
        .upload-icon { font-size: 48px; margin-bottom: 16px; }
        .upload-area h3 { font-size: 18px; font-weight: 700; color: #333; margin-bottom: 8px; }
        .upload-area p { font-size: 14px; color: #888; }
        #fileInput { display: none; }
        .file-selected { background: #f0faf4; border: 2px solid #2ecc71; border-radius: 16px; padding: 16px 24px; display: flex; align-items: center; gap: 12px; margin-top: 16px; }
        .file-selected span { font-size: 14px; font-weight: 600; color: #27ae60; }
        .analyze-btn { width: 100%; padding: 16px; background: linear-gradient(135deg, #2ecc71, #27ae60); color: white; border: none; border-radius: 14px; font-size: 16px; font-weight: 700; cursor: pointer; margin-top: 20px; font-family: inherit; box-shadow: 0 4px 16px rgba(46,204,113,0.4); }
        .analyze-btn:disabled { opacity: 0.6; cursor: not-allowed; }
        .loading { text-align: center; padding: 40px; display: none; }
        .loading p { font-size: 16px; color: #2ecc71; font-weight: 600; margin-top: 16px; }
        .spinner { width: 48px; height: 48px; border: 4px solid #d4f5e4; border-top-color: #2ecc71; border-radius: 50%; animation: spin 1s linear infinite; margin: 0 auto; }
        @keyframes spin { to { transform: rotate(360deg); } }
        .results { display: none; }
        .results-card { background: white; border-radius: 24px; padding: 40px; box-shadow: 0 4px 24px rgba(0,0,0,0.07); }
        .results-card h2 { font-size: 24px; font-weight: 800; color: #1a1a1a; margin-bottom: 24px; padding-bottom: 16px; border-bottom: 2px solid #f0faf4; }
        .feedback { font-size: 15px; line-height: 1.8; color: #333; white-space: pre-wrap; }
        .new-btn { padding: 12px 28px; background: transparent; border: 2px solid #2ecc71; border-radius: 14px; color: #2ecc71; font-size: 14px; font-weight: 700; cursor: pointer; margin-top: 24px; font-family: inherit; }
    </style>
</head>
<body>
<div class="navbar">
    <div class="logo">📄</div>
    <h1>AI Resume Analyzer · Internee.pk</h1>
</div>
<div class="container">
    <div class="hero">
        <h2>Analyze Your Resume with AI 🚀</h2>
        <p>Upload your resume and get instant AI-powered feedback on skills, structure, and improvements</p>
    </div>
    <div class="upload-card" id="uploadSection">
        <div class="upload-area" onclick="document.getElementById('fileInput').click()">
            <div class="upload-icon">📁</div>
            <h3>Click to upload your resume</h3>
            <p>Supports PDF, DOCX, DOC and TXT files</p>
        </div>
        <input type="file" id="fileInput" accept=".pdf,.docx,.doc,.txt" onchange="fileSelected(this)">
        <div class="file-selected" id="fileSelected" style="display:none">
            <span>📄</span>
            <span id="fileName">No file selected</span>
        </div>
        <button class="analyze-btn" id="analyzeBtn" onclick="analyzeResume()" disabled>
            🔍 Analyze Resume
        </button>
    </div>
    <div class="loading" id="loading">
        <div class="spinner"></div>
        <p>AI is analyzing your resume... Please wait ✨</p>
    </div>
    <div class="results" id="results">
        <div class="results-card">
            <h2>📊 Resume Analysis Results</h2>
            <div class="feedback" id="feedback"></div>
            <button class="new-btn" onclick="resetForm()">← Analyze Another Resume</button>
        </div>
    </div>
</div>
<script>
    function fileSelected(input) {
        if (input.files && input.files[0]) {
            document.getElementById('fileName').textContent = input.files[0].name;
            document.getElementById('fileSelected').style.display = 'flex';
            document.getElementById('analyzeBtn').disabled = false;
        }
    }
    async function analyzeResume() {
        const fileInput = document.getElementById('fileInput');
        if (!fileInput.files[0]) return;
        document.getElementById('uploadSection').style.display = 'none';
        document.getElementById('loading').style.display = 'block';
        document.getElementById('results').style.display = 'none';
        const formData = new FormData();
        formData.append('file', fileInput.files[0]);
        try {
            const response = await fetch('http://127.0.0.1:5000/analyze', {
                method: 'POST',
                body: formData
            });
            const data = await response.json();
            document.getElementById('loading').style.display = 'none';
            document.getElementById('results').style.display = 'block';
            document.getElementById('feedback').textContent = data.feedback || data.error;
        } catch(e) {
            console.error('Error:', e);
            alert('Something went wrong. Please try again.');
            resetForm();
        }
    }
    function resetForm() {
        document.getElementById('uploadSection').style.display = 'block';
        document.getElementById('loading').style.display = 'none';
        document.getElementById('results').style.display = 'none';
        document.getElementById('fileInput').value = '';
        document.getElementById('fileSelected').style.display = 'none';
        document.getElementById('analyzeBtn').disabled = true;
    }
</script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    filename = file.filename.lower()
    
    try:
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file)
        elif filename.endswith('.doc'):
            text = extract_text_from_docx(file)
        elif filename.endswith('.txt'):
            text = extract_text_from_txt(file)
        else:
            # Try to read as text anyway
            try:
                text = file.read().decode('utf-8', errors='ignore')
            except:
                return jsonify({'error': 'Unsupported file format. Please upload PDF, DOCX, DOC or TXT.'}), 400
        
        if not text.strip():
            return jsonify({'error': 'Could not extract text from file. Please try a different file.'}), 400
        
        feedback = analyze_resume(text)
        return jsonify({'feedback': feedback})
    
    except Exception as e:
        print("Error:", str(e))
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)